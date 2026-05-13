"""
Convert an Ultrathink audit markdown document to a Word .docx on Desktop.

Usage:
  python scripts/export-audit-to-docx.py <input.md>
  python scripts/export-audit-to-docx.py -            # read from stdin

Writes to:
  C:\\Users\\Administrator\\Desktop\\EXPO_PERFECTION_AUDIT_<YYYY-MM-DD>.docx

Handles: H1/H2/H3/H4, paragraphs, bullets (- / *), numbered lists (1.),
fenced code blocks (```), blockquotes (>), inline `code`, **bold**, *italic*,
horizontal rules (---), tables (pipe-separated).
"""
import sys
import re
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Walk a line and add runs for **bold**, *italic*, `code`."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('*') and token.endswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render(md_text, out_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                r = p.add_run('\n'.join(code_buf))
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif re.match(r'^[-*=_]{3,}\s*$', line):
            doc.add_paragraph('_' * 60)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(line[2:])
            r.italic = True
        elif re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, text)
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, text)
        elif '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[-:\s|]+\|?\s*$', lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            for idx, h in enumerate(header_cells):
                hdr[idx].text = h
                for run in hdr[idx].paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows, start=1):
                cells = table.rows[r_idx].cells
                for c_idx, val in enumerate(row[:len(header_cells)]):
                    cells[c_idx].text = val
            continue
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

        i += 1

    doc.save(out_path)
    print(f'Wrote: {out_path}')


def main():
    if len(sys.argv) < 2:
        print('usage: python scripts/export-audit-to-docx.py <input.md|->', file=sys.stderr)
        sys.exit(2)

    src = sys.argv[1]
    if src == '-':
        md_text = sys.stdin.read()
    else:
        with open(src, 'r', encoding='utf-8') as f:
            md_text = f.read()

    today = date.today().isoformat()
    desktop = os.path.join(os.environ.get('USERPROFILE', 'C:\\Users\\Administrator'), 'Desktop')
    out_name = f'EXPO_PERFECTION_AUDIT_{today}.docx'
    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
    else:
        out_path = os.path.join(desktop, out_name)

    render(md_text, out_path)


if __name__ == '__main__':
    main()
